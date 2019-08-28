from docx import Document

doc = Document("666.docx")

# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)
#         break
#     break
tables = doc.tables
table = tables[0]
# print(table)
# cell = table.cell(0, 0)
# print(cell.text)
# cell = table.cell(0, 1)
# print(cell.text)
# cell = table.cell(0, 2)
# print(cell.text)
# cell = table.cell(0, 3)
# print(cell.text)
# cell = table.cell(0, 4)
# print(cell.text)
# cell = table.cell(0, 5)
# print(cell.text)
# cell = table.cell(0, 6)
# print(cell.text)
# cell = table.cell(0, 7)
# print(cell.text)
# cell = table.cell(0, 8)
# print(cell.text)
# cell = table.cell(0, 9)
# print(cell.text)

hdr_cells = table.rows[0].cells
print(hdr_cells[0].text)
print(hdr_cells[1].text)
print(hdr_cells[2].text)
